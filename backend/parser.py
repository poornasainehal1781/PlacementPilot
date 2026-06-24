import re
import os
import spacy
from PyPDF2 import PdfReader
from docx import Document

# We will load en_core_web_sm, but we wrap it in a function that installs/downloads it if not present
_nlp = None

def get_nlp():
    global _nlp
    if _nlp is None:
        try:
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("[Parser] SpaCy model 'en_core_web_sm' not found. Installing...")
            import subprocess
            import sys
            subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
            _nlp = spacy.load("en_core_web_sm")
    return _nlp

# Standard skill list for keyword matching
COMMON_SKILLS = [
    # Programming Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "go", "golang", "rust", "php", "swift", "kotlin", "sql", "r", "html", "css", "bash",
    # Frameworks & Libraries
    "react", "angular", "vue", "next.js", "nuxt", "svelte", "flask", "django", "fastapi", "spring boot", "laravel", "express", "node.js", "node", "jquery", 
    "bootstrap", "tailwind", "pytorch", "tensorflow", "keras", "scikit-learn", "sklearn", "pandas", "numpy", "spacy", "nltk", "opencv",
    # Databases & Caching
    "mysql", "postgresql", "postgres", "mongodb", "sqlite", "redis", "elasticsearch", "cassandra", "mariadb", "dynamodb", "oracle",
    # DevOps & Cloud
    "docker", "kubernetes", "aws", "azure", "gcp", "google cloud", "git", "github", "gitlab", "jenkins", "ansible", "terraform", "ci/cd", "circleci", "heroku",
    # Analytics & AI
    "machine learning", "deep learning", "nlp", "natural language processing", "data science", "data analysis", "tableau", "power bi", "hadoop", "spark",
    # Methodologies & Concepts
    "agile", "scrum", "project management", "system design", "rest api", "graphql", "microservices", "oop", "mvc", "test driven development", "tdd"
]

CAPITALIZATION_MAPS = {
    "aws": "AWS",
    "gcp": "GCP",
    "sql": "SQL",
    "html": "HTML",
    "css": "CSS",
    "api": "API",
    "rest api": "REST API",
    "graphql": "GraphQL",
    "ci/cd": "CI/CD",
    "oop": "OOP",
    "mvc": "MVC",
    "tdd": "TDD",
    "nlp": "NLP",
    "mysql": "MySQL",
    "postgresql": "PostgreSQL",
    "sqlite": "SQLite",
    "mongodb": "MongoDB",
    "next.js": "Next.js",
    "nuxt": "Nuxt",
    "javascript": "JavaScript",
    "typescript": "TypeScript"
}

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"[Parser] Error parsing PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"[Parser] Error parsing DOCX {docx_path}: {e}")
    return text

def extract_text(file_path):
    ext = file_path.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext == "docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")

def extract_email(text):
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    match = re.search(email_pattern, text)
    if not match:
        return None
        
    email = match.group(0)
    parts = email.split('@')
    if len(parts) != 2:
        return email
    local, domain = parts
    
    # 1. Truncate domain after TLD (e.g. gmail.comhttps -> gmail.com)
    domain_match = re.match(r'^([a-zA-Z0-9-]+(?:\.[a-zA-Z0-9-]+)*\.(?:com|org|net|in|co|edu|gov|mil|io|co\.uk|info|me|us|ca|de|fr|jp|br|au))', domain, re.IGNORECASE)
    if domain_match:
        domain = domain_match.group(1)
        
    # 2. Clean the local part using words from the first line (often full name)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        first_line = lines[0].lower()
        name_words = re.findall(r'[a-z0-9]+', first_line)
        name_words.sort(key=len, reverse=True)
        
        for word in name_words:
            if len(word) >= 3 and word in local.lower():
                idx = local.lower().index(word)
                if idx > 0:
                    prefix = local[:idx].lower()
                    common_prefixes = ["pe", "e", "p", "m", "g", "l", "email", "mail", "envelope", "contact", "phone", "profile", "portfolio", "github", "linkedin"]
                    if prefix in common_prefixes:
                        local = local[idx:]
                        break
                        
    return f"{local.lower()}@{domain.lower()}"

def extract_phone(text):
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,5}\)?[-.\s]?\d{2,5}(?:[-.\s]?\d{4,5})?'
    for match in re.finditer(phone_pattern, text):
        val = match.group(0)
        digits_only = re.sub(r'\D', '', val)
        if len(digits_only) >= 9:
            return val
    return None

def extract_name(text):
    nlp = get_nlp()
    doc = nlp(text[:1000]) # Scan the first 1000 characters for name
    # First priority: SpaCy PERSON entities
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            # Clean up the name a bit (remove newlines, extra spaces)
            name = re.sub(r'\s+', ' ', name)
            if len(name.split()) >= 2: # Name usually has at least first and last name
                return name
    
    # Fallback: Assume the very first non-empty line is the name
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        return lines[0]
    return "Unknown Candidate"

def extract_skills(text):
    text_lower = text.lower()
    skills_found = []
    for skill in COMMON_SKILLS:
        # Use regex boundary matching to prevent partial word matches (e.g. 'c' inside 'cat')
        pattern = r'\b' + re.escape(skill) + r'\b'
        # Special check for C++ and C#
        if skill == "c++":
            pattern = r'c\+\+'
        elif skill == "c#":
            pattern = r'c\#'
            
        if re.search(pattern, text_lower):
            # Normalize display name using our mappings or title casing
            display_name = CAPITALIZATION_MAPS.get(skill.lower(), skill.title())
            skills_found.append(display_name)
    return list(set(skills_found))

def extract_education(text):
    education_keywords = ["bachelor", "master", "phd", "degree", "university", "college", "school", "b.s", "m.s", "b.tech", "m.tech", "mba", "b.a", "m.a"]
    lines = text.split("\n")
    edu_list = []
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        line_lower = line_clean.lower()
        if any(keyword in line_lower for keyword in education_keywords):
            if len(line_clean) < 120:  # Ignore too long paragraphs that might be a false positive
                edu_list.append(line_clean)
                
    return edu_list[:5] # Limit to top 5 items

def extract_experience(text):
    # Regex to find sections like Experience, Employment, Work History
    lines = text.split("\n")
    experience_keywords = ["experience", "work history", "employment history", "professional experience", "positions held"]
    
    exp_list = []
    is_exp_section = False
    
    # Simple heuristic to extract experience statements
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        line_lower = line_clean.lower()
        
        # Check if line marks start of experience section
        if any(ekey in line_lower for ekey in experience_keywords) and len(line_clean) < 30:
            is_exp_section = True
            continue
            
        # Check if line marks start of education or other section
        if is_exp_section and any(ekey in line_lower for ekey in ["education", "skills", "projects", "certifications", "interests"]) and len(line_clean) < 30:
            is_exp_section = False
            break
            
        if is_exp_section:
            # We add lines that look like company names, titles, or descriptions
            if len(line_clean) > 10:
                exp_list.append(line_clean)
                
    # If no section was structured, search for lines containing year patterns
    if not exp_list:
        year_pattern = r'(19|20)\d{2}\s*[-–—]\s*(Present|(19|20)\d{2})'
        for line in lines:
            if re.search(year_pattern, line):
                if len(line.strip()) < 150:
                    exp_list.append(line.strip())
                    
    return exp_list[:15] # Limit to top 15 items for parsed dashboard overview

def parse_resume(file_path):
    raw_text = extract_text(file_path)
    if not raw_text.strip():
        raise ValueError("Could not extract any text from the uploaded file.")
        
    return {
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "skills": extract_skills(raw_text),
        "experience": extract_experience(raw_text),
        "education": extract_education(raw_text),
        "raw_text": raw_text
    }
